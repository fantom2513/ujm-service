import io

from docx import Document

from app.services.files.docx import parse_docx


def _docx_with_table_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Intro paragraph")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Header A"
    table.rows[0].cells[1].text = "Header B"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def test_parse_docx_includes_table_text():
    # Regression: `.paragraphs` alone skips table content entirely, unlike
    # the TS original (mammoth.extractRawText), which walks the whole body.
    result = await parse_docx(_docx_with_table_bytes())
    assert "Intro paragraph" in result
    assert "Header A" in result
    assert "Header B" in result


async def test_parse_docx_returns_string_for_invalid_buffer():
    result = await parse_docx(b"not a docx")
    assert isinstance(result, str)


async def test_parse_docx_never_throws_on_garbage_input():
    result = await parse_docx(bytes([0x00, 0x01, 0x02, 0xFF, 0xFE]))
    assert isinstance(result, str)


async def test_parse_docx_returns_empty_string_on_empty_buffer():
    result = await parse_docx(b"")
    assert result == ""
